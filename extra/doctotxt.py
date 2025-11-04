import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os

def convert_docx_to_txt():
    """
    Convert a .docx file to .txt with GUI file selection.
    """
    # Create a root window (hidden)
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    # Ask user to select input .docx file
    print("Please select a .docx file to convert...")
    input_file = filedialog.askopenfilename(
        title="Select DOCX file to convert",
        filetypes=[
            ("Word Documents", "*.docx"),
            ("All Files", "*.*")
        ]
    )
    
    if not input_file:
        print("No file selected. Exiting...")
        return
    
    print(f"Selected file: {input_file}")
    
    try:
        # Read the .docx file
        doc = Document(input_file)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text_content = '\n'.join(full_text)
        
        # Ask user where to save the .txt file
        default_name = os.path.splitext(os.path.basename(input_file))[0] + ".txt"
        
        print("Please choose where to save the .txt file...")
        output_file = filedialog.asksaveasfilename(
            title="Save TXT file as",
            defaultextension=".txt",
            initialfile=default_name,
            filetypes=[
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        
        if not output_file:
            print("No output file selected. Exiting...")
            return
        
        # Write to .txt file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"\n✓ Conversion successful!")
        print(f"  Input:  {input_file}")
        print(f"  Output: {output_file}")
        
        # Show success message
        messagebox.showinfo(
            "Success", 
            f"File converted successfully!\n\nSaved to:\n{output_file}"
        )
        
    except Exception as e:
        error_msg = f"Error converting file: {str(e)}"
        print(f"\n✗ {error_msg}")
        messagebox.showerror("Error", error_msg)
    
    finally:
        root.destroy()


if __name__ == "__main__":
    print("=" * 60)
    print("DOCX to TXT Converter")
    print("=" * 60)
    convert_docx_to_txt()